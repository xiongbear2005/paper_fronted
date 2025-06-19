import mammoth
import tempfile
import os
from docx import Document
import base64
import re
from pathlib import Path
import io
import docx
from docx.document import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from services.docx2html import Docx2HtmlConverter

def convert_word_to_html(uploaded_file):
    """将 Word 文档转换为 HTML（基础版，不进行公式处理）"""
    try:
        # 读取上传的文件字节数据
        bytes_data = uploaded_file.getvalue()

        # 直接使用 mammoth 将 docx 转为 html
        result = mammoth.convert_to_html(io.BytesIO(bytes_data))
        html = result.value

        # 若检测到目录标记，则只保留从目录开始的内容，逻辑与旧版保持一致
        toc_markers = ["目录", "contents", "table of contents"]
        for marker in toc_markers:
            pattern = re.compile(f'<[^>]*>{re.escape(marker)}</[^>]*>', re.IGNORECASE)
            match = pattern.search(html)
            if match:
                html = html[match.start():]
                break

        # 最简单的样式包装，后续可再扩展
        styled_html = f"""
        <div style="font-family: 'Inter', system-ui, -apple-system, sans-serif; line-height: 1.6; color: var(--text-primary);">
            {html}
        </div>
        """

        return styled_html
    except Exception as e:
        print(f"转换Word文档时出错: {e}")
        return None

def convert_word_to_html_with_math(uploaded_file):
    """
    将 Word 文档转换为 HTML（增强版，支持公式、图片和复杂格式）
    使用docx2html.py进行转换，支持数学公式的渲染
    
    Args:
        uploaded_file: Streamlit上传的文件对象
        
    Returns:
        str: 生成的HTML内容
    """
    try:
        # 创建临时文件夹用于处理
        with tempfile.TemporaryDirectory() as temp_dir:
            # 保存上传的文件到临时目录
            temp_docx_path = os.path.join(temp_dir, "temp_document.docx")
            with open(temp_docx_path, "wb") as f:
                f.write(uploaded_file.getvalue())
            
            # 设置输出HTML路径
            temp_html_path = os.path.join(temp_dir, "temp_document.html")
            
            # 使用Docx2HtmlConverter进行转换
            converter = Docx2HtmlConverter()
            converter.convert_docx_to_html(
                docx_path=temp_docx_path,
                output_path=temp_html_path,
                title=uploaded_file.name,
                include_images=True
            )
            
            # 读取生成的HTML文件
            with open(temp_html_path, "r", encoding="utf-8") as f:
                html_content = f.read()
            
            # 读取生成的图片文件夹（如果存在）
            images_dir = os.path.join(temp_dir, "temp_document_images")
            image_files = {}
            if os.path.exists(images_dir):
                for img_file in os.listdir(images_dir):
                    img_path = os.path.join(images_dir, img_file)
                    with open(img_path, "rb") as img:
                        # 转换图片为base64编码，以便嵌入HTML
                        image_data = base64.b64encode(img.read()).decode('utf-8')
                        mime_type = get_mime_type(img_file)
                        image_files[img_file] = f"data:{mime_type};base64,{image_data}"
                
                # 替换HTML中的图片引用为base64编码
                for img_file, img_data in image_files.items():
                    img_dir_name = os.path.basename(images_dir)
                    img_path = f"{img_dir_name}/{img_file}"
                    html_content = html_content.replace(f'src="{img_path}"', f'src="{img_data}"')
            
            return html_content
    except Exception as e:
        print(f"使用增强版转换器处理Word文档时出错: {e}")
        # 如果增强版转换失败，回退到基础版
        return convert_word_to_html(uploaded_file)

def get_mime_type(file_path):
    """根据文件扩展名确定MIME类型"""
    ext = os.path.splitext(file_path)[1].lower()
    mime_types = {
        '.png': 'image/png',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.gif': 'image/gif',
        '.bmp': 'image/bmp',
        '.svg': 'image/svg+xml',
        '.webp': 'image/webp'
    }
    return mime_types.get(ext, 'image/png')  # 默认为PNG

def extract_toc_from_docx(uploaded_file):
    """从Word文档中提取目录结构，优化识别"第X章"式标题和子章节"""
    try:
        # 读取文档
        doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
        
        toc_items = []
        main_chapters = []  # 存储主章节
        sub_chapters = []   # 存储子章节
        
        # 标记用于处理目录和文档区域
        found_toc = False  # 是否找到"目录"
        in_toc_section = False  # 是否在目录区域内
        found_content_start = False  # 是否找到正文开始（第二次出现"第一章"）
        chapter_texts = set()  # 用于跟踪已添加的章节，防止重复
        
        print("开始分析文档结构...")
        
        # 第一次扫描：查找目录区域和正文开始位置
        toc_start_index = -1  # 目录开始位置
        toc_end_index = -1    # 目录结束位置
        content_start_index = -1  # 正文开始位置
        
        # 章节计数器，用于检测重复章节
        chapter_count = {}
        
        # 第一次扫描：查找目录和章节位置
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            # 检测目录节
            if not found_toc and (text.lower() == "目录" or text.lower() == "contents" or text.lower() == "table of contents"):
                found_toc = True
                in_toc_section = True
                toc_start_index = i
                print(f"检测到目录: '{text}' at index {i}")
                continue
            
            # 如果已经找到目录，检测章节标题
            if found_toc and in_toc_section:
                # 检测章节标题模式
                chapter_match = re.match(r'^第[一二三四五六七八九十\d]+章', text)
                if chapter_match:
                    chapter_name = chapter_match.group(0)
                    
                    # 记录章节出现次数
                    if chapter_name not in chapter_count:
                        chapter_count[chapter_name] = 1
                    else:
                        chapter_count[chapter_name] += 1
                    
                    # 如果是第二次出现"第一章"，表示正文开始
                    if chapter_name in ["第一章", "第1章"] and chapter_count.get(chapter_name, 0) > 1:
                        in_toc_section = False
                        found_content_start = True
                        content_start_index = i
                        toc_end_index = i - 1  # 上一段落是目录的结束
                        print(f"检测到正文开始: '{text}' at index {i}")
                        break
            
        # 如果没有明确找到目录结束位置，但找到了正文开始，则以正文开始位置作为目录结束
        if toc_end_index == -1 and content_start_index != -1:
            toc_end_index = content_start_index - 1
            
        print(f"文档分析结果: 目录开始={toc_start_index}, 目录结束={toc_end_index}, 正文开始={content_start_index}")
        
        # 如果找到了目录区域，提取章节结构
        if found_toc and toc_start_index != -1 and content_start_index != -1:
            # 从正文开始位置提取章节
            start_index = content_start_index
            
            # 第二次扫描：从正文开始位置提取章节
            for i, paragraph in enumerate(doc.paragraphs[start_index:], start_index):
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # 识别为标题样式的段落
                is_heading = False
                level = 0
                
                # 通过样式名识别标题
                if paragraph.style.name.startswith('Heading') or '标题' in paragraph.style.name:
                    is_heading = True
                    try:
                        # 从样式名获取级别
                        level_match = re.search(r'\d+', paragraph.style.name)
                        if level_match:
                            level = int(level_match.group(0))
                        else:
                            level = 1  # 默认为一级标题
                    except:
                        level = 1
                
                # 通过格式识别标题 - 检查是否粗体或大字体
                elif paragraph.runs:
                    is_bold = any(run.bold for run in paragraph.runs if hasattr(run, 'bold'))
                    is_large = False
                    try:
                        is_large = any(run.font.size and run.font.size.pt > 14 for run in paragraph.runs if hasattr(run, 'font') and hasattr(run.font, 'size'))
                    except:
                        pass
                        
                    if is_bold or is_large:
                        # 进一步检查是否匹配章节标题模式
                        is_heading = True
                        level = 1 if is_large else 2
                
                # 通过文本模式识别标题
                chapter_patterns = [
                    # 主章节模式
                    (r'^第[一二三四五六七八九十\d]+章[：:]?\s*\S+', 1),  # 第一章：绪论
                    (r'^第[一二三四五六七八九十\d]+章$', 1),  # 第一章
                    (r'^\d+[\.、]\s*[^\.、\d]+', 1),  # 1. 绪论
                    (r'^[一二三四五六七八九十]+[\.、]\s*[^\.、\d]+', 1),  # 一. 绪论
                    
                    # 子章节模式
                    (r'^\d+\.\d+\s+\S+', 2),  # 1.1 研究背景
                    (r'^\d+\.\d+\.\d+\s+\S+', 3),  # 1.1.1 具体内容
                    (r'^第[一二三四五六七八九十\d]+节\s+\S+', 2),  # 第一节 内容
                ]
                
                if not is_heading:
                    for pattern, pat_level in chapter_patterns:
                        if re.match(pattern, text):
                            is_heading = True
                            level = pat_level
                            break
                
                # 标准化章节名称（移除数字后缀等）
                standardized_text = standardize_chapter_name(text)
                display_text = standardized_text[:20] + "..." if len(standardized_text) > 20 else standardized_text
                original_text = text  # 保留原始文本用于匹配
                
                # 章节标题需满足以下额外条件：
                # 1. 不能包含中文或英文标点（如"：，。.?!等"）
                # 2. 一级章节标题通常字体较大或明确使用 Heading 1 / 标题 1 样式

                # 过滤包含标点符号的标题（目录行通常带有省略号或页码）
                has_punctuation = re.search(r'[，。：；、,\.;:!？?!]', original_text) is not None

                # 检查字体大小是否足够大（>14pt 视为大字体）
                is_large_font = False
                try:
                    if paragraph.runs:
                        for run in paragraph.runs:
                            if run.font.size and run.font.size.pt and run.font.size.pt > 14:
                                is_large_font = True
                                break
                except Exception:
                    pass

                # 检查是否为一级标题样式（Heading 1 或 标题 1）
                style_name_lower = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
                is_heading1_style = style_name_lower.startswith('heading 1') or '标题 1' in paragraph.style.name or '标题1' in paragraph.style.name if paragraph.style else False

                # 对于 level==1，需要字体较大或使用 Heading 1 样式
                meets_font_style_requirement = True
                if level == 1:
                    meets_font_style_requirement = is_large_font or is_heading1_style

                # 跳过个人信息、不符合章节特征的内容、重复的章节、含标点或字体/样式不符合要求的标题
                if (
                    is_heading
                    and not is_personal_info(text)
                    and display_text not in chapter_texts
                    and not has_punctuation
                    and meets_font_style_requirement
                ):
                    # 添加到已处理章节集合，防止重复
                    chapter_texts.add(display_text)
                    
                    # 创建章节ID，结合段落索引和章节文本的前几个字符（确保唯一性）
                    safe_text = re.sub(r'[^\w\d]', '', text[:5])
                    chapter_id = f"section-{i}-{safe_text}"
                    
                    item = {
                        'index': i,
                        'level': level,
                        'text': display_text,
                        'original_text': original_text,  # 保存原始文本用于内容匹配
                        'standardized_text': standardized_text,  # 保存标准化后的文本
                        'id': chapter_id,
                        'children': []
                    }
                    
                    print(f"添加章节: level={level}, text='{display_text}', id={chapter_id}")
                    
                    if level == 1:
                        main_chapters.append(item)
                    else:
                        sub_chapters.append(item)
        
        # 如果没有找到足够的章节，尝试其他规则
        if len(main_chapters) < 1:
            print("未找到足够主章节，尝试通过内容特征识别...")
            
            # 查找有明显特征的段落
            start_index = content_start_index if content_start_index != -1 else 0
            for i, paragraph in enumerate(doc.paragraphs[start_index:start_index+100], start_index):
                text = paragraph.text.strip()
                if not text or len(text) < 4 or len(text) > 100:
                    continue
                
                # 标准化章节名称
                standardized_text = standardize_chapter_name(text)
                display_text = standardized_text[:20] + "..." if len(standardized_text) > 20 else standardized_text
                
                # 检查是否已添加
                if display_text in chapter_texts:
                    continue
                
                # 常见章节关键词
                chapter_keywords = ['绪论', '引言', '简介', '概述', '背景', '方法', '实验', '结果', '分析', 
                                   '讨论', '结论', '参考文献', '致谢', 'Introduction', 'Methods', 
                                   'Results', 'Discussion', 'Conclusion', 'References']
                
                if any(keyword in text for keyword in chapter_keywords):
                    # 确认文本格式特征 - 粗体或单独成段落等
                    is_formatted = False
                    if paragraph.runs:
                        is_bold = any(run.bold for run in paragraph.runs if hasattr(run, 'bold'))
                        if is_bold:
                            is_formatted = True
                    
                    # 如果是单独的短段落也可能是标题
                    if len(text) < 30:
                        is_formatted = True
                    
                    if is_formatted and not is_personal_info(text):
                        # 添加到已处理章节集合
                        chapter_texts.add(display_text)
                        
                        # 创建章节ID，结合段落索引和章节文本的前几个字符
                        safe_text = re.sub(r'[^\w\d]', '', text[:5])
                        chapter_id = f"section-{i}-{safe_text}"
                        
                        main_chapters.append({
                            'index': i,
                            'level': 1,
                            'text': display_text,
                            'original_text': text,
                            'standardized_text': standardized_text,
                            'id': chapter_id,
                            'children': []
                        })
                        
                        print(f"通过内容特征添加章节: '{display_text}', id={chapter_id}")
        
        # 构建层级关系
        for sub_item in sub_chapters:
            # 找到最近的主章节作为父级
            parent_found = False
            for i in range(len(main_chapters)-1, -1, -1):
                if main_chapters[i]['index'] < sub_item['index']:
                    main_chapters[i]['children'].append(sub_item)
                    parent_found = True
                    break
            
            # 如果没找到父级，可能是独立的子章节或序言等
            if not parent_found and main_chapters:
                main_chapters[0]['children'].append(sub_item)
        
        # 合并所有章节数据
        toc_items = main_chapters
        
        print(f"共提取 {len(toc_items)} 个主章节, {len(sub_chapters)} 个子章节")
        for i, chapter in enumerate(toc_items):
            print(f"  {i+1}. {chapter['text']} (ID: {chapter['id']})")
            for j, subchapter in enumerate(chapter.get('children', [])):
                print(f"     {i+1}.{j+1} {subchapter['text']} (ID: {subchapter['id']})")
        
        return toc_items
    except Exception as e:
        print(f"提取目录时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return []

def standardize_chapter_name(text):
    """标准化章节名称，去除数字后缀等"""
    # 处理"第X章 绪论1"这样的情况，转换为"第X章 绪论"
    text = re.sub(r'^(第[一二三四五六七八九十\d]+章\s*[\s\S]*绪论)\d+', r'\1', text)
    text = re.sub(r'^(第[一二三四五六七八九十\d]+章\s*[\s\S]*引言)\d+', r'\1', text)
    
    # 去除标题末尾的数字
    text = re.sub(r'(第[一二三四五六七八九十\d]+[章节].*?[^0-9])\d+$', r'\1', text)
    
    # 清理额外的空格
    text = re.sub(r'\s+', ' ', text).strip()
    
    return text

def is_likely_content_start(paragraphs, current_text):
    """判断是否是目录结束、正文开始的位置"""
    # 检查前面的段落是否有明显的格式变化
    format_change = False
    content_indicators = False
    
    # 检查格式变化（例如从短段落到长段落）
    text_lengths = [len(p.text.strip()) for p in paragraphs if p.text.strip()]
    if text_lengths and (max(text_lengths) < 30) and len(current_text) > 30:
        format_change = True
    
    # 检查是否有页眉、页脚、分隔符等内容指示
    for p in paragraphs:
        text = p.text.strip().lower()
        if text and (len(text) < 5 or text in ['page', 'chapter', '-'*3]):
            content_indicators = True
            break
    
    # 检查当前文本是否是明显的章节开始
    is_chapter_start = re.match(r'^第[一二三四五六七八九十\d]+章\s+\S+', current_text) is not None
    
    return format_change or content_indicators or is_chapter_start

def is_personal_info(text):
    """检查文本是否包含个人信息"""
    personal_info_patterns = [
        r'姓名[：:]\s*\w+',
        r'电话[：:]\s*\d+',
        r'邮箱[：:]\s*[\w\.-]+@[\w\.-]+',
        r'地址[：:]\s*\w+',
        r'学号[：:]\s*\w+',
        r'指导教师[：:]\s*\w+'
    ]
    
    for pattern in personal_info_patterns:
        if re.search(pattern, text):
            return True
    return False

def simulate_analysis_with_toc(uploaded_file):
    """模拟文档分析，整合目录提取和内容转换"""
    try:
        # 提取目录结构
        toc_items = extract_toc_from_docx(uploaded_file)
        
        # 仅转换为 HTML（不再生成 Markdown）
        html_content = convert_word_to_html(uploaded_file)
        
        # 模拟分析结果
        word_count = sum(len(paragraph.text.split()) for paragraph in docx.Document(io.BytesIO(uploaded_file.getvalue())).paragraphs if paragraph.text.strip())
        
        # 模拟特殊关键词提取
        special_keywords = ["研究", "分析", "方法", "结果", "讨论"]
        
        # 生成论文整体总结评价
        paper_summary = {
            "overall_comment": "本论文结构清晰，论述有理有据，研究方法得当，实验设计合理，结果可靠。整体来看，论文在选题、方法和数据分析方面表现出了较高的学术水平。",
            "strengths": [
                "研究问题明确，具有重要的理论和现实意义",
                "研究方法选择恰当，技术路线清晰可行",
                "数据分析全面，结果呈现清晰直观",
                "结论基于实证分析，具有较强的说服力"
            ],
            "weaknesses": [
                "文献综述部分对最新研究的涵盖不够全面",
                "研究方法部分的理论依据可进一步加强",
                "数据分析中的一些假设条件需要更充分的说明",
                "对研究局限性的讨论可以更加深入"
            ],
            "suggestions": [
                "建议补充最近1-2年发表的相关领域最新研究成果",
                "可以增加对研究方法选择的理论依据说明",
                "建议在数据分析部分增加敏感性分析，以增强结果的稳健性",
                "可以在结论部分更明确地指出未来研究方向"
            ]
        }
        
        # 为每个章节生成分析结果
        for chapter in toc_items:
            # 添加分析字段
            chapter['analysis'] = {
                'summary': f"本章节主要讨论{chapter.get('standardized_text', chapter.get('text', ''))}相关内容，包含了相关理论基础和研究方法。",
                'strengths': [
                    "结构清晰，层次分明",
                    "论述有理有据，引用恰当",
                    "关键概念解释详细"
                ],
                'weaknesses': [
                    f"建议为{chapter.get('standardized_text', chapter.get('text', ''))}添加更详细的小节划分，以提高内容的组织性",
                    f"考虑在{chapter.get('standardized_text', chapter.get('text', ''))}中添加具体案例或数据支持，增强说服力",
                    "部分论述可以更加精炼，避免冗余"
                ]
            }
            
            # 如果有子章节，添加相关建议
            if 'children' in chapter and chapter['children']:
                chapter['analysis']['subchapter_advice'] = f"建议加强{chapter.get('standardized_text', chapter.get('text', ''))}与其{len(chapter['children'])}个子章节之间的过渡说明，使内容衔接更加自然流畅。"
        
        # 构建分析结果
        analysis_result = {
            'word_count': word_count,
            'special_keywords': special_keywords,
            'chapters': toc_items,
            'paper_summary': paper_summary,
            'html_content': html_content
        }
        
        return analysis_result
    except Exception as e:
        print(f"分析文档时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return None 